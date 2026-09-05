"""Generate legal Word documents matching law report publication standards:
- Narrow margins (0.5 in / 36 pt on all sides)
- Subject / Catchwords Title at top (Book Antiqua 12 B, Center)
- Journal header (Book Antiqua 10 Italic, Center)
- Standardized Court Name (e.g. BOMBAY HIGH COURT) & Bench (e.g. (Aurangabad Bench))
- Coram directly under Bench (e.g. (ABHAY S. WAGHWASE, J.))
- Parties in Title Case with '- Applicant' and '- Respondents'
- Bordered case details block (Appeal number, connected applications, Decided on date)
- HeadNote block (Book Antiqua 11 B, Spacing Before: 7 pt, After: 4 pt)
- Law Point block (Book Antiqua 12 B, Spacing Before: 3.6 pt, After: 4.3 pt)
- List of Cases Referred (Heading: Minion Pro 11.5 BU; Cases: Minion Pro 11.5, Left: 0.4 cm)
- Advocates Appeared for the Parties (Book Antiqua 10.5 Italic, Left: 0.21 cm)
- Clean JUDGMENT heading (no duplicate judge names or labels)
- Decimal numbered paragraphs: 0 tab spacing, Para Number: Book Antiqua 11 B
- Sub-paragraphs strictly split and indented with 1 tab spacing (0.5 inch / first-line indent)
- Blockquotes indented (Left: 1.0 in, Right: 0.5 in, Italic)
- ORDER centered with 0 indent, and Roman numeral order items with 1 tab spacing (0.5 in)
"""
from __future__ import annotations

from io import BytesIO
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from parser_ai import JudgmentSchema, JudgmentBlock, has_devanagari


def _apply_font(run, font_name: str, size: float, bold: bool = False, italic: bool = False, underline: bool = False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)


def add_top_and_bottom_borders(paragraph):
    """Add clean single borders to top and bottom of the case details box."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom'):
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '6')
        bdr.set(qn('w:space'), '6')
        bdr.set(qn('w:color'), 'auto')
        pBdr.append(bdr)
    pPr.append(pBdr)


def generate_docx(judgment: JudgmentSchema, font_size: int = 12) -> BytesIO:
    doc = Document()

    # Narrow margins (0.5 in / 36 pt on all sides)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 1. Subject / Catchwords Title
    subj_title = getattr(judgment, 'subject_title', '')
    if subj_title:
        p_subj = doc.add_paragraph()
        p_subj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_subj.paragraph_format.space_before = Pt(0)
        p_subj.paragraph_format.space_after = Pt(2)
        _apply_font(p_subj.add_run(subj_title.strip()), "Book Antiqua", font_size, bold=True)

    # 2. Journal Header Block
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.paragraph_format.space_after = Pt(2)
    _apply_font(p_hdr.add_run(getattr(judgment, 'journal_header', '') or "(2026 Maharashtra e Journal)"), "Book Antiqua", 10, italic=True)

    # 3. Court Name
    court = getattr(judgment, 'court_name', '')
    if court:
        p_court = doc.add_paragraph()
        p_court.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_court.paragraph_format.space_after = Pt(2)
        _apply_font(p_court.add_run(court.upper()), "Book Antiqua", font_size, bold=True)

    # 4. Bench
    bench = getattr(judgment, 'bench', '')
    if bench:
        p_bench = doc.add_paragraph()
        p_bench.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_bench.paragraph_format.space_after = Pt(2)
        bench_text = bench.strip()
        if not bench_text.startswith("("):
            bench_text = f"({bench_text})"
        _apply_font(p_bench.add_run(bench_text), "Book Antiqua", font_size, bold=True)

    # 5. Coram (Placed directly below Bench)
    coram = getattr(judgment, 'coram', '')
    if coram:
        p_coram = doc.add_paragraph()
        p_coram.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_coram.paragraph_format.space_after = Pt(6)
        coram_text = coram.strip()
        if not coram_text.startswith("("):
            coram_text = f"({coram_text})"
        _apply_font(p_coram.add_run(coram_text), "Book Antiqua", font_size, bold=True)

    # 6. Parties Block
    appellant = getattr(judgment, 'appellant', '')
    if appellant:
        p_app = doc.add_paragraph()
        p_app.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_app.paragraph_format.space_after = Pt(2)
        app_role = getattr(judgment, 'appellant_role', '') or "Applicant"
        _apply_font(p_app.add_run(f"{appellant.strip()} - {app_role}"), "Book Antiqua", font_size, bold=True)

        p_vs = doc.add_paragraph()
        p_vs.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_vs.paragraph_format.space_after = Pt(2)
        _apply_font(p_vs.add_run("Versus"), "Book Antiqua", font_size, italic=True)

    respondent = getattr(judgment, 'respondent', '')
    if respondent:
        p_resp = doc.add_paragraph()
        p_resp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_resp.paragraph_format.space_after = Pt(6)
        resp_role = getattr(judgment, 'respondent_role', '') or "Respondents"
        _apply_font(p_resp.add_run(f"{respondent.strip()} - {resp_role}"), "Book Antiqua", font_size, bold=True)

    # 7. Case Details Block (Bordered)
    details_content = getattr(judgment, 'case_details', '').strip()
    if not details_content:
        parts = [p for p in [getattr(judgment, 'appeal_number', ''), getattr(judgment, 'judgment_date', '')] if p]
        details_content = "   |   ".join(parts)

    if details_content:
        p_details = doc.add_paragraph()
        p_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_details.paragraph_format.space_before = Pt(3.6)
        p_details.paragraph_format.space_after = Pt(2.2)
        _apply_font(p_details.add_run(details_content), "Book Antiqua", 10.5, bold=False)
        add_top_and_bottom_borders(p_details)

    # 8. HeadNote
    headnotes = getattr(judgment, 'headnotes', [])
    if headnotes:
        for hn in headnotes:
            is_mar = has_devanagari(hn)
            p_hn = doc.add_paragraph()
            p_hn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if is_mar:
                p_hn.paragraph_format.space_after = Pt(10)
                _apply_font(p_hn.add_run(hn.strip()), "Tiro Devanagari Marathi", 12)
            else:
                p_hn.paragraph_format.space_before = Pt(7)
                p_hn.paragraph_format.space_after = Pt(4)
                _apply_font(p_hn.add_run(hn.strip()), "Book Antiqua", 11, bold=True)

    # 9. Law Point
    law_points = getattr(judgment, 'law_points', [])
    if law_points:
        p_lp_head = doc.add_paragraph()
        p_lp_head.paragraph_format.space_before = Pt(3.6)
        p_lp_head.paragraph_format.space_after = Pt(2)
        _apply_font(p_lp_head.add_run("Law Point"), "Book Antiqua", font_size, bold=True)

        for lp in law_points:
            is_mar = has_devanagari(lp)
            p_lp = doc.add_paragraph()
            p_lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if is_mar:
                p_lp.paragraph_format.space_before = Pt(3.6)
                p_lp.paragraph_format.space_after = Pt(4.3)
                _apply_font(p_lp.add_run(lp.strip()), "Tiro Devanagari Marathi", 11, bold=True)
            else:
                p_lp.paragraph_format.space_before = Pt(2)
                p_lp.paragraph_format.space_after = Pt(4.3)
                _apply_font(p_lp.add_run(lp.strip()), "Book Antiqua", font_size, bold=True)

    # 10. List of Cases Referred
    cases_referred = getattr(judgment, 'cases_referred', [])
    if cases_referred:
        p_cr_head = doc.add_paragraph()
        p_cr_head.paragraph_format.space_before = Pt(6)
        p_cr_head.paragraph_format.space_after = Pt(2)
        _apply_font(p_cr_head.add_run("List of Cases Referred :"), "Minion Pro", 11.5, bold=True, underline=True)

        for i, case in enumerate(cases_referred):
            clean_case = re.sub(r"^\d+\.\s*", "", case.strip())
            p_case = doc.add_paragraph()
            p_case.paragraph_format.left_indent = Cm(0.4)
            p_case.paragraph_format.space_before = Pt(2)
            p_case.paragraph_format.space_after = Pt(2)
            _apply_font(p_case.add_run(f"{i+1}. {clean_case}"), "Minion Pro", 11.5)

    # 11. Advocates Block (Supports both advocates_block and advocates list)
    adv_text = getattr(judgment, 'advocates_block', '').strip()
    if not adv_text:
        adv_list = getattr(judgment, 'advocates', [])
        if adv_list:
            adv_text = "Advocates Appeared for the Parties :- " + "; ".join(adv_list) + ";"

    if adv_text:
        p_adv = doc.add_paragraph()
        p_adv.paragraph_format.left_indent = Cm(0.21)
        p_adv.paragraph_format.space_before = Pt(5)
        p_adv.paragraph_format.space_after = Pt(4.1)

        prefix = "Advocates Appeared for the Parties :- "
        if adv_text.startswith(prefix):
            _apply_font(p_adv.add_run(prefix), "Book Antiqua", 10.5, bold=True, italic=True)
            _apply_font(p_adv.add_run(adv_text[len(prefix):]), "Book Antiqua", 10.5, italic=True)
        else:
            _apply_font(p_adv.add_run(prefix), "Book Antiqua", 10.5, bold=True, italic=True)
            _apply_font(p_adv.add_run(adv_text), "Book Antiqua", 10.5, italic=True)

    # 12. JUDGMENT Heading (Clean, single heading)
    p_judg = doc.add_paragraph()
    p_judg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_judg.paragraph_format.space_before = Pt(3.6)
    p_judg.paragraph_format.space_after = Pt(3.6)
    _apply_font(p_judg.add_run("JUDGMENT"), "Book Antiqua", font_size, bold=True)

    # 13. Judgment Body Paragraphs
    body_blocks = getattr(judgment, 'judgment_body', [])
    for block in body_blocks:
        text = block.text.strip()
        if not text:
            continue

        # Prevent duplicate headers
        if re.fullmatch(r"JUDGMENT\s*:?", text, re.I) or re.match(r"^ABHAY S\. WAGHWASE.*", text, re.I):
            continue

        font_name = "Tiro Devanagari Marathi" if block.is_marathi else "Book Antiqua"

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15

        # Standalone ORDER heading
        if re.fullmatch(r"ORDER", text, re.I) or block.type == "order":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _apply_font(p.add_run(text), font_name, font_size, bold=True)
            continue

        # Divider
        if block.type == "divider" or text == "-----------------------------":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _apply_font(p.add_run("-----------------------------"), font_name, font_size)
            continue

        # Order items (I), (II), (III) -> 1 tab spacing (0.5 inch)
        if block.type == "order_item" or re.match(r"^\([IVXLCDMivxlcdm]+\)", text):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(5)
            _apply_font(p.add_run(text), font_name, font_size)
            continue

        # Blockquotes (indented 1.0 inch left, 0.5 inch right)
        if block.type == "blockquote" or text.startswith(('"', '“', '‘')):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(1.0)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(5)
            _apply_font(p.add_run(text), font_name, font_size - 0.5, italic=True)
            continue

        # Decimal numbered paragraphs (1., 2., ...) -> 0 tab spacing, Para Number 11 B
        dec_match = re.match(r"^(\d+(?:\.\d+)*[\.\)]\s*)(.*)", text, re.DOTALL)
        if dec_match and block.type != "subparagraph":
            prefix, body_text = dec_match.group(1), dec_match.group(2)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(5)
            _apply_font(p.add_run(prefix), font_name, 11, bold=True)
            if body_text:
                _apply_font(p.add_run(body_text), font_name, font_size)
            continue

        # Sub-paragraphs -> 1 tab spacing (0.5 inch / first-line indent)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(5)
        _apply_font(p.add_run(text), font_name, font_size)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output